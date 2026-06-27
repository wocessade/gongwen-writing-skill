#!/usr/bin/env python3
"""Generate a document processing slip (发文处理笺) in .docx format.

Usage:
    python create_processing_slip.py --slip-type "发文处理笺" --source-org "XX省人民政府" --doc-number "X政发〔2024〕1号" --date "2024年12月20日" --title "标题" --circulation "张三,李四,王五" --output "./output/处理笺.docx"
    python create_processing_slip.py --config path/to/config.json --output "./output/处理笺.docx"
"""

import json, os, sys, warnings
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ── Font Helpers ──

_FONT_FALLBACKS = {
    '方正小标宋简体': ('宋体', '方正小标宋简体 not installed; falling back to 宋体.'),
    '仿宋_GB2312':   ('仿宋', '仿宋_GB2312 not installed; falling back to 仿宋.'),
}
_FONT_SAFE = {'宋体', '仿宋', '黑体', '楷体', 'SimSun', 'SimHei', 'FangSong', 'KaiTi', 'Times New Roman'}

def resolve_font(font_cn):
    if font_cn in _FONT_SAFE:
        return font_cn
    if font_cn in _FONT_FALLBACKS:
        fallback, msg = _FONT_FALLBACKS[font_cn]
        warnings.warn(msg, stacklevel=2)
        return fallback
    return font_cn

def set_run_font(run, font_cn, font_en='宋体', size=Pt(16), bold=False):
    run.font.size = size
    run.bold = bold
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)

def set_cell_border(cell, **kwargs):
    """Set cell border. Usage: set_cell_border(cell, top={"sz":4,"color":"000000","val":"single"}, ...)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, attrs in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        for attr, val in attrs.items():
            element.set(qn(f'w:{attr}'), str(val))
        tcBorders.append(element)
    tcPr.append(tcBorders)

# ── Document Creation ──

def create_processing_slip(slip_type="发文处理笺", source_org="", doc_number="",
                           date="", title="", circulation=None):
    """Create a document processing slip docx.

    Args:
        slip_type: "发文处理笺" or "收文处理笺"
        source_org: Source organ name
        doc_number: Document number
        date: Date string
        title: Document title
        circulation: List of names for circulation

    Returns:
        python-docx Document object
    """
    if circulation is None:
        circulation = []

    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Title ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(slip_type)
    set_run_font(title_run, resolve_font('方正小标宋简体'), font_en='SimSun', size=Pt(22), bold=True)

    # ── Metadata table ──
    meta_fields = [
        ("来文机关", source_org),
        ("文  号", doc_number),
        ("收文日期", date),
        ("标  题", title),
    ]

    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table width
    tbl = meta_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    for i, (label, value) in enumerate(meta_fields):
        # Label cell (left, ~20% width)
        label_cell = meta_table.cell(i, 0)
        label_cell.width = Cm(2.5)
        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(label)
        set_run_font(label_run, resolve_font('仿宋_GB2312'), size=Pt(14), bold=True)
        set_cell_border(label_cell, top={"sz":"4","color":"000000","val":"single"},
                        bottom={"sz":"4","color":"000000","val":"single"},
                        left={"sz":"4","color":"000000","val":"single"},
                        right={"sz":"4","color":"000000","val":"single"})

        # Value cell (right, ~80% width)
        val_cell = meta_table.cell(i, 1)
        val_cell.width = Cm(13.0)
        val_para = val_cell.paragraphs[0]
        val_run = val_para.add_run(value)
        set_run_font(val_run, resolve_font('仿宋_GB2312'), size=Pt(14))
        set_cell_border(val_cell, top={"sz":"4","color":"000000","val":"single"},
                        bottom={"sz":"4","color":"000000","val":"single"},
                        left={"sz":"4","color":"000000","val":"single"},
                        right={"sz":"4","color":"000000","val":"single"})

    # ── 拟办意见 ──
    doc.add_paragraph()  # spacer
    opinion_table = doc.add_table(rows=6, cols=2)
    opinion_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = opinion_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    # Merge all right cells vertically
    right_cells = [opinion_table.cell(r, 1) for r in range(6)]
    if len(right_cells) > 1:
        right_cells[0].merge(right_cells[-1])

    # Label in first cell
    label_cell = opinion_table.cell(0, 0)
    label_cell.width = Cm(2.5)
    label_para = label_cell.paragraphs[0]
    label_run = label_para.add_run("拟办意见")
    set_run_font(label_run, resolve_font('仿宋_GB2312'), size=Pt(14), bold=True)
    set_cell_border(label_cell, top={"sz":"4","color":"000000","val":"single"},
                    bottom={"sz":"4","color":"000000","val":"single"},
                    left={"sz":"4","color":"000000","val":"single"},
                    right={"sz":"4","color":"000000","val":"single"})

    # Set height for merged right cell
    merged_right = right_cells[0]
    set_cell_border(merged_right, top={"sz":"4","color":"000000","val":"single"},
                    bottom={"sz":"4","color":"000000","val":"single"},
                    left={"sz":"4","color":"000000","val":"single"},
                    right={"sz":"4","color":"000000","val":"single"})

    # ── 批示意见 ──
    doc.add_paragraph()  # spacer
    approval_table = doc.add_table(rows=6, cols=2)
    approval_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = approval_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    right_cells2 = [approval_table.cell(r, 1) for r in range(6)]
    if len(right_cells2) > 1:
        right_cells2[0].merge(right_cells2[-1])

    label_cell2 = approval_table.cell(0, 0)
    label_cell2.width = Cm(2.5)
    label_para2 = label_cell2.paragraphs[0]
    label_run2 = label_para2.add_run("批示意见")
    set_run_font(label_run2, resolve_font('仿宋_GB2312'), size=Pt(14), bold=True)
    set_cell_border(label_cell2, top={"sz":"4","color":"000000","val":"single"},
                    bottom={"sz":"4","color":"000000","val":"single"},
                    left={"sz":"4","color":"000000","val":"single"},
                    right={"sz":"4","color":"000000","val":"single"})

    merged_right2 = right_cells2[0]
    set_cell_border(merged_right2, top={"sz":"4","color":"000000","val":"single"},
                    bottom={"sz":"4","color":"000000","val":"single"},
                    left={"sz":"4","color":"000000","val":"single"},
                    right={"sz":"4","color":"000000","val":"single"})

    # ── 传阅签字表 ──
    doc.add_paragraph()  # spacer
    num_rows = max(len(circulation) + 1, 4)  # header + circulation + at least 2 empty
    sign_table = doc.add_table(rows=num_rows, cols=4)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = sign_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    headers = ["序号", "姓名", "签字", "日期"]
    col_widths = [Cm(1.5), Cm(3.0), Cm(5.5), Cm(5.5)]

    for col_idx, (hdr, width) in enumerate(zip(headers, col_widths)):
        cell = sign_table.cell(0, col_idx)
        cell.width = width
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(hdr)
        set_run_font(run, resolve_font('仿宋_GB2312'), size=Pt(14), bold=True)
        set_cell_border(cell, top={"sz":"4","color":"000000","val":"single"},
                        bottom={"sz":"4","color":"000000","val":"single"},
                        left={"sz":"4","color":"000000","val":"single"},
                        right={"sz":"4","color":"000000","val":"single"})

    for i, name in enumerate(circulation):
        row_idx = i + 1
        values = [str(row_idx), name, "", ""]
        for col_idx, (val, width) in enumerate(zip(values, col_widths)):
            cell = sign_table.cell(row_idx, col_idx)
            cell.width = width
            para = cell.paragraphs[0]
            if col_idx <= 1:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(val)
            set_run_font(run, resolve_font('仿宋_GB2312'), size=Pt(14))
            set_cell_border(cell, top={"sz":"4","color":"000000","val":"single"},
                            bottom={"sz":"4","color":"000000","val":"single"},
                            left={"sz":"4","color":"000000","val":"single"},
                            right={"sz":"4","color":"000000","val":"single"})

    # Empty rows for additional signatures
    for i in range(len(circulation) + 1, num_rows):
        for col_idx in range(4):
            cell = sign_table.cell(i, col_idx)
            para = cell.paragraphs[0]
            set_cell_border(cell, top={"sz":"4","color":"000000","val":"single"},
                            bottom={"sz":"4","color":"000000","val":"single"},
                            left={"sz":"4","color":"000000","val":"single"},
                            right={"sz":"4","color":"000000","val":"single"})

    return doc


def load_config_from_json(path):
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def main():
    import argparse
    parser = argparse.ArgumentParser(
        description='Generate a document processing slip (发文/收文处理笺) in .docx format.'
    )
    parser.add_argument('--slip-type', default='发文处理笺', choices=['发文处理笺', '收文处理笺'],
                        help='Type of processing slip (default: 发文处理笺)')
    parser.add_argument('--source-org', default='', help='来文机关')
    parser.add_argument('--doc-number', default='', help='文号')
    parser.add_argument('--date', default='', help='收文日期')
    parser.add_argument('--title', default='', help='公文标题')
    parser.add_argument('--circulation', default='', help='传阅人员，逗号分隔')
    parser.add_argument('--config', type=str, help='Path to JSON config file')
    parser.add_argument('--output', '-o', default='', help='Output .docx path')
    args = parser.parse_args()

    # If config file provided, load from it
    if args.config:
        config = load_config_from_json(args.config)
    else:
        config = {
            'slip_type': args.slip_type,
            'source_org': args.source_org,
            'doc_number': args.doc_number,
            'date': args.date,
            'title': args.title,
            'circulation': [n.strip() for n in args.circulation.split(',') if n.strip()],
        }

    doc = create_processing_slip(
        slip_type=config.get('slip_type', '发文处理笺'),
        source_org=config.get('source_org', ''),
        doc_number=config.get('doc_number', ''),
        date=config.get('date', ''),
        title=config.get('title', ''),
        circulation=config.get('circulation', []),
    )

    output_path = args.output or config.get('output', '')
    if not output_path:
        safe = config.get('doc_number', '处理笺').replace('/', '_').replace('\\', '_')
        output_path = os.path.join(os.getcwd(), f'{safe}.docx')

    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)

    print(f"[OK] Processing slip saved to: {output_path}")
    print(f"     Tables: {len(doc.tables)}")

    return output_path


if __name__ == '__main__':
    main()
