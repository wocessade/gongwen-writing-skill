#!/usr/bin/env python3
"""Convert existing .docx documents to GB/T 9704-2012 compliant format.

Analyzes document structure (headings, body) by content patterns,
then restyles paragraphs to conform to the Chinese government
document formatting standard.

Usage:
    python convert_gongwen.py --input nonstandard.docx --output compliant.docx
    python convert_gongwen.py --input-dir ./input/ --output-dir ./output/ --verbose
    python convert_gongwen.py -i nonstandard.docx --dry-run -v
"""

import os
import re
import sys
import warnings
import argparse
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Font Helpers ──

_FONT_FALLBACKS = {
    '方正小标宋简体': ('宋体', '方正小标宋简体 not installed; falling back to 宋体. Install the font for compliant red headers.'),
    '仿宋_GB2312':   ('仿宋', '仿宋_GB2312 not installed; falling back to 仿宋. Install the font for compliant body text.'),
    '楷体_GB2312':   ('楷体', '楷体_GB2312 not installed; falling back to 楷体. Install the font for compliant level-2 headings.'),
}
_FONT_SAFE = {'宋体', '仿宋', '黑体', '楷体', 'SimSun', 'SimHei', 'FangSong', 'KaiTi', 'Times New Roman'}


def resolve_font(font_cn):
    if font_cn in _FONT_SAFE:
        return font_cn
    if font_cn in _FONT_FALLBACKS:
        fallback, msg = _FONT_FALLBACKS[font_cn]
        warnings.warn(msg, stacklevel=2)
        return fallback
    return font_cn


def set_run_font(run, font_cn, font_en='宋体', size=Pt(16), bold=False):
    run.font.size = size
    run.bold = bold
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)


# ── Pattern Definitions ──

_TITLE_PATTERN = re.compile(r'关于.+的(通知|报告|函|请示|纪要|决定|通报|批复|意见|通告|公告|命令|令|议案|决议|公报)')
_HEADING0_PATTERN = re.compile(r'^[一二三四五六七八九十]+、')
_HEADING1_PATTERN = re.compile(r'^（[一二三四五六七八九十]+）')


def _classify_paragraph(text):
    """Classify a paragraph by its text content.

    Returns one of: 'title', 'heading0', 'heading1', 'body'
    """
    text = text.strip()
    if not text:
        return 'body'
    if _TITLE_PATTERN.search(text):
        return 'title'
    if _HEADING0_PATTERN.match(text):
        return 'heading0'
    if _HEADING1_PATTERN.match(text):
        return 'heading1'
    return 'body'


def analyze_document(doc):
    """Analyze current formatting of a document.

    Returns:
        dict with:
        - margins: dict of page margins
        - total_paragraphs: int
        - paragraphs: list of per-para analysis dicts
    """
    section = doc.sections[0]
    margins = {
        'top_margin': round(section.top_margin.cm, 2) if section.top_margin else None,
        'bottom_margin': round(section.bottom_margin.cm, 2) if section.bottom_margin else None,
        'left_margin': round(section.left_margin.cm, 2) if section.left_margin else None,
        'right_margin': round(section.right_margin.cm, 2) if section.right_margin else None,
    }

    paras = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        info = {
            'index': i,
            'text': text[:60] + ('...' if len(text) > 60 else ''),
            'class': _classify_paragraph(text),
            'alignment': str(para.alignment) if para.alignment else 'None',
        }
        # Check first run for font info
        if para.runs:
            r = para.runs[0]
            info['font_name'] = r.font.name
            info['font_size'] = str(r.font.size) if r.font.size else None
            info['bold'] = r.bold
        paras.append(info)

    return {
        'margins': margins,
        'total_paragraphs': len(doc.paragraphs),
        'paragraphs': paras,
    }


def restyle_paragraph(para, para_class):
    """Restyle a paragraph to GB/T 9704-2012 standards.

    Args:
        para: python-docx Paragraph object.
        para_class: One of 'title', 'heading0', 'heading1', 'body'.
    """
    pf = para.paragraph_format

    # Reset paragraph formatting
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(28)

    if para_class == 'title':
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Cm(0)
        font_name = resolve_font('方正小标宋简体')
        font_size = Pt(22)
        bold = False
    elif para_class == 'heading0':
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
        font_name = resolve_font('黑体')
        font_size = Pt(16)
        bold = True
    elif para_class == 'heading1':
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
        font_name = resolve_font('楷体_GB2312')
        font_size = Pt(16)
        bold = False
    else:  # body
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0.85)
        font_name = resolve_font('仿宋_GB2312')
        font_size = Pt(16)
        bold = False

    for run in para.runs:
        set_run_font(run, font_name, size=font_size, bold=bold)


def convert_document(input_path, output_path=None, dry_run=False, verbose=False):
    """Convert a single .docx to GB/T 9704-2012 compliant format.

    Args:
        input_path: Path to input .docx.
        output_path: Path to output .docx. If None, generates from input name.
        dry_run: If True, only analyze and print, don't save.
        verbose: If True, print per-paragraph analysis.

    Returns:
        Output path if not dry_run, else None.
    """
    input_path = os.path.abspath(input_path)
    if not os.path.isfile(input_path):
        print(f"[ERROR] File not found: {input_path}", file=sys.stderr)
        return None

    print(f"[INFO] Analyzing: {input_path}")
    doc = Document(input_path)

    analysis = analyze_document(doc)

    # Print analysis
    print(f"       Page margins: top={analysis['margins']['top_margin']}cm "
          f"bottom={analysis['margins']['bottom_margin']}cm "
          f"left={analysis['margins']['left_margin']}cm "
          f"right={analysis['margins']['right_margin']}cm")
    print(f"       Total paragraphs: {analysis['total_paragraphs']}")

    if verbose:
        print(f"\n{'Idx':>4} {'Class':<10} {'Text'}")
        print("-" * 60)
        for p in analysis['paragraphs']:
            print(f"{p['index']:>4} {p['class']:<10} {p['text']}")

    if dry_run:
        print(f"[DRY-RUN] No changes saved.")
        return None

    # Restyle each paragraph
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        para_class = _classify_paragraph(text)
        restyle_paragraph(para, para_class)

    # Fix page margins to GB standard
    section = doc.sections[0]
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    # Save
    if not output_path:
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_converted{ext}"

    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)

    file_size = os.path.getsize(output_path)
    print(f"[OK] Converted document saved to: {output_path}")
    print(f"     File size: {file_size:,} bytes")

    return output_path


def main():
    parser = argparse.ArgumentParser(
        description='Convert existing .docx documents to GB/T 9704-2012 compliant format.'
    )
    parser.add_argument('--input', '-i', type=str,
                        help='Input .docx file path.')
    parser.add_argument('--input-dir', '-I', type=str,
                        help='Directory of .docx files to batch convert.')
    parser.add_argument('--output', '-o', type=str, default=None,
                        help='Output .docx path (for single file mode).')
    parser.add_argument('--output-dir', '-O', type=str, default=None,
                        help='Output directory (for batch mode).')
    parser.add_argument('--dry-run', '-n', action='store_true',
                        help='Show analysis without saving changes.')
    parser.add_argument('--verbose', '-v', action='store_true',
                        help='Print per-paragraph analysis.')
    args = parser.parse_args()

    if not args.input and not args.input_dir:
        parser.print_help()
        print("\n[ERROR] Provide --input or --input-dir.", file=sys.stderr)
        sys.exit(1)

    if args.input:
        convert_document(args.input, args.output, args.dry_run, args.verbose)

    if args.input_dir:
        input_dir = os.path.abspath(args.input_dir)
        if not os.path.isdir(input_dir):
            print(f"[ERROR] Input directory not found: {input_dir}", file=sys.stderr)
            sys.exit(1)

        output_dir = os.path.abspath(args.output_dir) if args.output_dir else f"{input_dir}_converted"
        os.makedirs(output_dir, exist_ok=True)

        docx_files = [f for f in sorted(os.listdir(input_dir))
                      if f.lower().endswith('.docx')]
        if not docx_files:
            print(f"[WARN] No .docx files found in {input_dir}")
            return

        print(f"[INFO] Found {len(docx_files)} .docx files in {input_dir}")
        success = 0
        for fname in docx_files:
            in_path = os.path.join(input_dir, fname)
            out_path = os.path.join(output_dir, fname)
            result = convert_document(in_path, out_path, args.dry_run, args.verbose)
            if result:
                success += 1

        print(f"\n[DONE] Batch convert: {success}/{len(docx_files)} successful")


if __name__ == '__main__':
    main()
