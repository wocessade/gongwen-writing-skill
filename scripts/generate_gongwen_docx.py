#!/usr/bin/env python3
"""Generate a Chinese government official document (.docx) conforming to GB/T 9704-2012.

党政机关公文格式生成脚本
Reference: GB/T 9704-2012 党政机关公文格式

Usage:
    python generate_gongwen_docx.py --config path/to/config.json
    python generate_gongwen_docx.py  # generates a sample 通知 document
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import os
import sys
import warnings
import argparse


# ── Font Registry ──────────────────────────────────────────────────────────

# Font availability check: known to be problematic on stock Windows
_FONT_FALLBACKS = {
    '方正小标宋简体': ('宋体', '方正小标宋简体 not installed; falling back to 宋体. Install the font for compliant red headers.'),
    '仿宋_GB2312':   ('仿宋', '仿宋_GB2312 not installed; falling back to 仿宋. Install the font for compliant body text.'),
    '楷体_GB2312':   ('楷体', '楷体_GB2312 not installed; falling back to 楷体. Install the font for compliant level-2 headings.'),
}

# Fonts available on nearly all Chinese Windows systems
_FONT_SAFE = {'宋体', '仿宋', '黑体', '楷体', 'SimSun', 'SimHei', 'FangSong', 'KaiTi', 'Times New Roman'}


def resolve_font(font_cn):
    """Return the best available font, issuing a warning for missing special fonts.

    Args:
        font_cn: Desired Chinese font name.

    Returns:
        str: Font name to use (original or fallback).
    """
    if font_cn in _FONT_SAFE:
        return font_cn
    if font_cn in _FONT_FALLBACKS:
        fallback, msg = _FONT_FALLBACKS[font_cn]
        warnings.warn(msg, stacklevel=2)
        return fallback
    return font_cn


# ── Font Helpers ───────────────────────────────────────────────────────────

def set_run_font(run, font_cn, font_en='宋体', size=Pt(16), bold=False):
    """Set both East-Asian (eastAsia) and Latin (ascii + hAnsi) fonts on a run.

    This is the core formatting helper for all GB/T 9704-2012 text elements.
    Uses OxmlElement for w:rFonts to ensure both font axes are set correctly.

    Args:
        run: python-docx Run object.
        font_cn: Chinese font name for w:eastAsia attribute.
        font_en: Latin font name for w:ascii and w:hAnsi (default 'Times New Roman').
        size: Font size in Pt (default Pt(16) = 三号).
        bold: Bold flag (default False).
    """
    run.font.size = size
    run.bold = bold
    # Built-in heading styles default to blue in Word; always reset to black.
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)


# ── Page Setup ─────────────────────────────────────────────────────────────

def set_page_setup(section):
    """Configure A4 page margins per GB/T 9704-2012 section 5.1.

    - Page: A4 (210mm x 297mm)
    - Top margin (天头): 37mm ± 1mm
    - Bottom margin (地脚): 35mm ± 1mm
    - Left margin (订口): 28mm ± 1mm
    - Right margin (翻口): 26mm ± 1mm
    - 版心尺寸: 156mm x 225mm

    Args:
        section: python-docx Section object.
    """
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)


# ── Content Elements ───────────────────────────────────────────────────────

def _red_header_size(name: str) -> Pt:
    """Calculate red header font size based on authority name length.

    GB/T 9704-2012 §7.2.4: 发文机关标志最大高度 ≤ 22mm ≈ 62pt.
    No hardcoded size is specified — agencies determine based on
    principle of visual prominence (醒目美观).

    This heuristic sizes by character count (including spaces), aiming
    to keep the full name within 版心宽度 156mm:
      2-3 chars → Pt(55)   (e.g. 国务院, 教育部)
      4-5 chars → Pt(48)   (e.g. XX省人民政府)
      6-7 chars → Pt(42)   (e.g. XX省人民政府办公厅)
      8-9 chars → Pt(36)   (e.g. XX市人力资源和社会保障局)
      10+ chars → Pt(30)   (联合行文 long names)
    """
    n = len(name)
    if n <= 3:
        return Pt(55)
    elif n <= 5:
        return Pt(48)
    elif n <= 7:
        return Pt(42)
    elif n <= 9:
        return Pt(36)
    else:
        return Pt(30)


def add_red_header(doc, authority_name):
    """Add the red document header (发文机关标志) per GB/T 9704-2012 section 7.2.4.

    - Red (RGB 255, 0, 0) 方正小标宋简体 (with fallback to 宋体)
    - Centered
    - Maximum height <= 22mm (dynamic sizing based on name length)
    - Followed by a red horizontal divider line (红色反线), ~0.44mm thick

    Args:
        doc: python-docx Document.
        authority_name: Full name of issuing authority, e.g. 'XX省人民政府'.
    """
    header_font = resolve_font('方正小标宋简体')
    font_size = _red_header_size(authority_name)
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0

    run = para.add_run(authority_name)
    set_run_font(run, header_font, font_en='SimSun', size=font_size, bold=False)
    run.font.color.rgb = RGBColor(255, 0, 0)

    # Red horizontal divider line (bottom border on the paragraph)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '10')     # 1.25pt ~ 0.44mm
    bottom.set(qn('w:color'), 'FF0000')
    bottom.set(qn('w:space'), '4')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_joint_red_header(doc, authorities):
    """联合行文红头：各发文机关分行排列。
    GB/T 9704-2012 §7.2.4: 联合行文时，各发文机关标志分行排列。
    每行按自身长度独立定字号，不互相钳制。
    """
    header_font = resolve_font('方正小标宋简体')
    for idx, name in enumerate(authorities):
        font_size = _red_header_size(name)
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.0
        run = para.add_run(name)
        set_run_font(run, header_font, font_en='SimSun', size=font_size, bold=False)
        run.font.color.rgb = RGBColor(255, 0, 0)

        # 红色分隔线（仅加在最后一个发文机关段落）
        if idx == len(authorities) - 1:
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '10')
            bottom.set(qn('w:color'), 'FF0000')
            bottom.set(qn('w:space'), '4')
            pBdr.append(bottom)
            pPr.append(pBdr)


def add_security(doc, security_text):
    """Add security classification (密级) per GB/T 9704-2012 section 7.2.1.

    黑体 三号(16pt) bold, left-aligned.

    Args:
        doc: python-docx Document.
        security_text: e.g. '绝密·永久', '机密★10年', '秘密'.
    """
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = para.add_run(security_text)
    set_run_font(run, resolve_font('黑体'), size=Pt(16), bold=True)


def add_urgency(doc, urgency_text):
    """Add urgency level (紧急程度) per GB/T 9704-2012 section 7.2.2.

    黑体 三号(16pt) bold, left-aligned.

    Args:
        doc: python-docx Document.
        urgency_text: e.g. '特急' or '加急'.
    """
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = para.add_run(urgency_text)
    set_run_font(run, resolve_font('黑体'), size=Pt(16), bold=True)


def add_document_number(doc, number_text):
    """Add document number (发文字号) per GB/T 9704-2012 section 7.2.5.

    仿宋_GB2312 三号(16pt), centered below red header divider line.

    Args:
        doc: python-docx Document.
        number_text: e.g. 'X政发〔2024〕1号'.
    """
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    run = para.add_run(number_text)
    set_run_font(run, resolve_font('仿宋_GB2312'), size=Pt(16))


def add_title(doc, text):
    """Add document title (标题) per GB/T 9704-2012 section 7.3.1.

    方正小标宋简体 二号(22pt), centered. Spacer blank line before title
    to maintain visual separation from the red header area.

    Args:
        doc: python-docx Document.
        text: Document title, e.g. '关于做好2024年元旦春节期间有关工作的通知'.
    """
    # Spacer line before title
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(28)

    title_font = resolve_font('方正小标宋简体')
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(28)
    run = para.add_run(text)
    set_run_font(run, title_font, font_en='SimSun', size=Pt(22))


def add_main_recipient(doc, text):
    """Add main recipient (主送机关) per GB/T 9704-2012 section 7.3.2.

    仿宋_GB2312 三号(16pt), left-aligned, ending with a colon.

    Args:
        doc: python-docx Document.
        text: Recipient name(s), e.g. '各省、自治区、直辖市人民政府：'.
    """
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(28)
    run = para.add_run(text)
    set_run_font(run, resolve_font('仿宋_GB2312'), size=Pt(16))


def add_body_para(doc, text, indent=True):
    """Add a body paragraph (正文) per GB/T 9704-2012 section 7.3.3.

    仿宋_GB2312 三号(16pt), 28pt fixed line spacing.
    每面22行, 每行28字.

    Args:
        doc: python-docx Document.
        text: Paragraph text.
        indent: Apply first-line indent of 2 characters (~0.85cm) if True.

    Returns:
        Paragraph: The added paragraph object.
    """
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(28)
    if indent:
        pf.first_line_indent = Cm(0.85)
    run = para.add_run(text)
    set_run_font(run, resolve_font('仿宋_GB2312'), size=Pt(16))
    return para


def add_heading_level1(doc, text):
    """Add a level-1 heading (一级标题) per GB/T 9704-2012 section 7.3.3.

    黑体 三号(16pt) bold, format '一、二、三、'.

    Args:
        doc: python-docx Document.
        text: Full heading text, e.g. '一、提高思想认识'.
    """
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(28)
    run = para.add_run(text)
    set_run_font(run, resolve_font('黑体'), size=Pt(16), bold=True)


def add_heading_level2(doc, text):
    """Add a level-2 heading (二级标题) per GB/T 9704-2012 section 7.3.3.

    楷体_GB2312 三号(16pt), format '（一）（二）（三）'.

    Args:
        doc: python-docx Document.
        text: Full heading text, e.g. '（一）充分认识做好元旦春节期间工作的重要性'.
    """
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(28)
    run = para.add_run(text)
    set_run_font(run, resolve_font('楷体_GB2312'), size=Pt(16))


def add_attachment(doc, attachments):
    """Add attachment description (附件说明) per GB/T 9704-2012 section 7.3.4.

    仿宋_GB2312 三号(16pt), placed after body text and before signature.

    Supports both flat string lists and nested attachments:
        attachments = [
            "XXX文件",
            {"title": "YYY文件", "children": ["YYY细则", "YYY附表"]},
        ]

    Args:
        doc: python-docx Document.
        attachments: List of attachment descriptions (str or dict).
    """
    if not attachments:
        return

    def _add_para(text, indent=Cm(0.85)):
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = Pt(28)
        pf.first_line_indent = indent
        run = para.add_run(text)
        set_run_font(run, resolve_font('仿宋_GB2312'), size=Pt(16))

    def _render(items, prefix='附件：', indent_level=0, counter=1):
        base_indent = Cm(0.85 + indent_level * 1.0)
        for item in items:
            if isinstance(item, str):
                _add_para(f'{prefix}{counter}. {item}', base_indent)
                counter += 1
            elif isinstance(item, dict):
                title = item.get('title', '')
                children = item.get('children', [])
                _add_para(f'{prefix}{counter}. {title}', base_indent)
                counter += 1
                if children:
                    _render(children, '附件：', indent_level + 1, 1)

    _render(attachments)


def add_signature_and_date(doc, org_name, date_str):
    """Add issuing organ signature and date (发文机关署名 + 成文日期) per GB/T 9704-2012 section 7.3.5.

    Right-aligned with 右空四字 (approximately 4-character right indent).
    Separated from body/attachments by a blank line.

    Args:
        doc: python-docx Document.
        org_name: Issuing organ name, e.g. 'XX省人民政府'.
        date_str: Date string, e.g. '2024年1月15日'.
    """
    # Spacer blank line before signature
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(28)

    # Issuing organ name
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(28)
    pf.right_indent = Cm(2.0)
    run = para.add_run(org_name)
    set_run_font(run, resolve_font('仿宋_GB2312'), size=Pt(16))

    # Date
    para2 = doc.add_paragraph()
    pf2 = para2.paragraph_format
    pf2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf2.space_before = Pt(0)
    pf2.space_after = Pt(0)
    pf2.line_spacing = Pt(28)
    pf2.right_indent = Cm(2.0)
    run2 = para2.add_run(date_str)
    set_run_font(run2, resolve_font('仿宋_GB2312'), size=Pt(16))


def add_joint_signature_and_date(doc, org_names, date_str):
    """联合行文署名：各机关分行排列右对齐，最后一个署名下压成文日期。"""
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(28)

    for name in org_names:
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = Pt(28)
        pf.right_indent = Cm(2.0)
        run = para.add_run(name)
        set_run_font(run, resolve_font('仿宋_GB2312'), size=Pt(16))

    para2 = doc.add_paragraph()
    pf2 = para2.paragraph_format
    pf2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf2.space_before = Pt(0)
    pf2.space_after = Pt(0)
    pf2.line_spacing = Pt(28)
    pf2.right_indent = Cm(2.0)
    run2 = para2.add_run(date_str)
    set_run_font(run2, resolve_font('仿宋_GB2312'), size=Pt(16))


def add_banji(doc, cc_list=None, printing_org=None, printing_date=None, copies=None):
    """Add complete 版记 (page notes) per GB/T 9704-2012 §7.4.

    Structure:
      [thick line]
      抄送：XXX (if cc_list)
      [thin line]
      印发机关              印发日期  (if both)
      共印X份                       (if copies)

    Args:
        doc: python-docx Document.
        cc_list: list of CC organ names.
        printing_org: 印发机关 name.
        printing_date: 印发日期 string (e.g. '2024年12月20日印发').
        copies: total copies number (int or str).
    """
    has_cc = bool(cc_list)
    has_print = bool(printing_org) or bool(printing_date)

    if not has_cc and not has_print:
        return

    def add_hline(sz='12'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(14)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), sz)
        top.set(qn('w:color'), '000000')
        top.set(qn('w:space'), '1')
        pBdr.append(top)
        pPr.append(pBdr)

    add_hline('12')  # 首条粗线
    if has_cc:
        cc_text = '抄送：' + '、'.join(cc_list)
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.left_indent = Cm(0.85)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = Pt(14)
        run = p.add_run(cc_text)
        set_run_font(run, resolve_font('仿宋_GB2312'), size=Pt(16))

    if has_print:
        add_hline('6')  # 中间细线
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = Pt(14)
        pf.left_indent = Cm(0.85)
        if printing_org:
            r1 = p.add_run(printing_org)
            set_run_font(r1, resolve_font('仿宋_GB2312'), size=Pt(16))
        if printing_date:
            r2 = p.add_run(f'\t\t{printing_date}' if printing_org else printing_date)
            set_run_font(r2, resolve_font('仿宋_GB2312'), size=Pt(16))

        if copies:
            p2 = doc.add_paragraph()
            pf2 = p2.paragraph_format
            pf2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pf2.right_indent = Cm(0.5)
            pf2.space_before = Pt(0)
            pf2.space_after = Pt(0)
            pf2.line_spacing = Pt(14)
            r3 = p2.add_run(f'共印{copies}份')
            set_run_font(r3, resolve_font('仿宋_GB2312'), size=Pt(16))

    add_hline('12')  # 末条粗线


def add_signatory(doc, signatory_name):
    """Add signatory (签发人) for 上行文 per GB/T 9704-2012 section 7.2.6.

    仿宋_GB2312 三号(16pt), placed to the right of the document number area.

    Args:
        doc: python-docx Document.
        signatory_name: Name of the approving signatory.
    """
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(28)
    run = para.add_run(f'签发人：{signatory_name}')
    set_run_font(run, resolve_font('仿宋_GB2312'), size=Pt(16))


# ── Page Number ────────────────────────────────────────────────────────────

def add_page_number(section):
    """Add page numbers in gongwen format per GB/T 9704-2012 section 7.5.

    宋体 四号(14pt), '- 1 -' format (一人线 around the page number).
    Single pages right-aligned, double pages left-aligned.

    Args:
        section: python-docx Section object.
    """
    # Even page footer — left-aligned
    even_footer = section.even_page_footer
    even_footer.is_linked_to_previous = False
    even_p = even_footer.paragraphs[0]
    even_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Odd page footer — right-aligned
    odd_footer = section.footer
    odd_footer.is_linked_to_previous = False
    odd_p = odd_footer.paragraphs[0]
    odd_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for p in [even_p, odd_p]:
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        # Left dash: —
        run_dash1 = p.add_run('— ')
        set_run_font(run_dash1, resolve_font('宋体'), font_en='宋体', size=Pt(14))

        # PAGE field (begin)
        run_begin = p.add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        run_begin._r.append(fldChar_begin)

        # instrText for PAGE
        run_instr = p.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run_instr._r.append(instrText)

        # Separate
        run_sep = p.add_run()
        fldChar_sep = OxmlElement('w:fldChar')
        fldChar_sep.set(qn('w:fldCharType'), 'separate')
        run_sep._r.append(fldChar_sep)

        # Page number placeholder (visible before F9 update)
        run_page = p.add_run('1')
        set_run_font(run_page, resolve_font('宋体'), font_en='宋体', size=Pt(14))

        # End
        run_end = p.add_run()
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run_end._r.append(fldChar_end)

        # Right dash: —
        run_dash2 = p.add_run(' —')
        set_run_font(run_dash2, resolve_font('宋体'), font_en='宋体', size=Pt(14))


def add_meeting_info(doc, meeting_info):
    """Add meeting information block for 纪要 format.

    Displays time, location, chair, attendees, recorder as labeled lines.
    仿宋_GB2312 三号(16pt), left-aligned with first-line indent.

    Args:
        doc: python-docx Document.
        meeting_info: dict with keys: time, location, chair, attendees, recorder.
            Attendees can be a list or string.
    """
    info_font = resolve_font('仿宋_GB2312')
    info_size = Pt(16)

    mapping = [
        ('时 间', 'time'),
        ('地 点', 'location'),
        ('主持人', 'chair'),
        ('参加人员', 'attendees'),
        ('记 录', 'recorder'),
    ]

    for label, key in mapping:
        val = meeting_info.get(key)
        if not val:
            continue
        if isinstance(val, (list, tuple)):
            val = '、'.join(str(v) for v in val)
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = Pt(28)
        pf.first_line_indent = Cm(0.85)
        run = para.add_run(f'{label}：{val}')
        set_run_font(run, info_font, size=info_size)


# ── Main Generator ─────────────────────────────────────────────────────────

def generate_gongwen(config):
    """Generate a GB/T 9704-2012 compliant Chinese government official document.

    This is the main entry point for the gongwen docx assembly pipeline.
    Call from the G2 strategist with a fully populated config dict.

    Config fields:
        authority (str): *Conditional*. 发文机关全称 (required if not joint).
        document_type (str): Required. 文种, used for filename generation.
        document_number (str): Required. 发文字号, e.g. 'X政发〔2024〕1号'.
        title (str): Required. 公文标题.
        main_recipient (str): Required. 主送机关.
        body_sections (list[dict]): Required. List of body content dicts.
            Each dict:
            - 'level': 0 (一级标题), 1 (二级标题), or None (正文段落)
            - 'text': Content string.
        issuing_org (str): *Conditional*. 发文机关署名 (required if not joint).
        date (str): Required. 成文日期, e.g. '2024年1月15日'.
        attachments (list[str], optional): 附件说明 items.
        cc_organs (list[str], optional): 抄送机关 list.
        security (str, optional): 密级, e.g. '机密'.
        urgency (str, optional): 紧急程度, e.g. '特急'.
        signatory (str, optional): 签发人, for 上行文.
        joint_authorities (list[str], optional): 联合行文发文机关列表（与authority二选一）.
        joint_issuing_orgs (list[str], optional): 联合行文署名机关列表（与issuing_org二选一）.
        printing_org (str, optional): 印发机关.
        printing_date (str, optional): 印发日期.
        copies (str|int, optional): 共印份数.
        meeting_info (dict, optional): 纪要信息，含 time/location/chair/attendees/recorder.
        output_path (str, optional): Output file path. Auto-generated if omitted.

    Returns:
        str: Absolute path to the generated .docx file.

    Raises:
        ValueError: If a required config field is missing.
    """
    # ── Validate required fields ──
    is_minutes = config.get('document_type') == '纪要'
    required = ['document_type', 'title', 'body_sections', 'date']
    if not is_minutes:
        required.append('document_number')
        required.append('main_recipient')
    for field in required:
        if field not in config or config[field] is None:
            raise ValueError(f"Missing required config field: '{field}'")

    is_joint = bool(config.get('joint_authorities'))
    is_joint_sig = bool(config.get('joint_issuing_orgs'))

    if not is_joint and ('authority' not in config or config['authority'] is None):
        raise ValueError("Must provide 'authority' (single) or 'joint_authorities' (joint)")
    if not is_joint_sig and ('issuing_org' not in config or config['issuing_org'] is None):
        raise ValueError("Must provide 'issuing_org' (single) or 'joint_issuing_orgs' (joint)")

    # ── Create document ──
    doc = Document()

    # Default style normalization
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(16)
    try:
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    except Exception:
        pass

    # ── Page Setup (section 7.1) ──
    section = doc.sections[0]
    set_page_setup(section)

    # ── Security classification (section 7.2.1) ──
    if config.get('security'):
        add_security(doc, config['security'])

    # ── Urgency (section 7.2.2) ──
    if config.get('urgency'):
        add_urgency(doc, config['urgency'])

    # ── Red header + divider line (section 7.2.4) — skip for 纪要 ──
    if is_minutes:
        pass  # 纪要 format: no red header
    elif is_joint:
        add_joint_red_header(doc, config['joint_authorities'])
    else:
        add_red_header(doc, config['authority'])

    # ── Document number (section 7.2.5) — skip for 纪要 ──
    if not is_minutes:
        add_document_number(doc, config['document_number'])

    # ── Signatory for 上行文 (section 7.2.6) — skip for 纪要 ──
    if config.get('signatory') and not is_minutes:
        add_signatory(doc, config['signatory'])

    # ── Title (section 7.3.1) ──
    add_title(doc, config['title'])

    # ── Meeting info (纪要 format) ──
    if is_minutes and config.get('meeting_info'):
        add_meeting_info(doc, config['meeting_info'])

    # ── Main recipient (section 7.3.2) — optional for 纪要 ──
    if config.get('main_recipient'):
        add_main_recipient(doc, config['main_recipient'])

    # ── Body sections (section 7.3.3) ──
    for item in config.get('body_sections', []):
        level = item.get('level')
        text = item.get('text', '')
        if not text:
            continue
        if level == 0:
            add_heading_level1(doc, text)
        elif level == 1:
            add_heading_level2(doc, text)
        else:
            add_body_para(doc, text, indent=True)

    # ── Attachments (section 7.3.4) ──
    add_attachment(doc, config.get('attachments', []))

    # ── Issuing organ signature + date (section 7.3.5) ──
    if is_joint_sig:
        add_joint_signature_and_date(doc, config['joint_issuing_orgs'], config['date'])
    else:
        add_signature_and_date(doc, config['issuing_org'], config['date'])

    # ── 版记 (section 7.4) ──
    add_banji(
        doc,
        cc_list=config.get('cc_organs', []),
        printing_org=config.get('printing_org'),
        printing_date=config.get('printing_date'),
        copies=config.get('copies'),
    )

    # ── Page numbers (section 7.5) ──
    add_page_number(section)

    # ── Save ──
    output_path = config.get('output_path')
    if not output_path:
        safe_type = config['document_type'].replace('/', '_')
        if is_joint:
            safe = config['joint_authorities'][0].replace('/', '_').replace('\\', '_')
            output_path = os.path.join(os.getcwd(), f'{safe}_等_{safe_type}.docx')
        else:
            safe = config['authority'].replace('/', '_').replace('\\', '_')
            output_path = os.path.join(os.getcwd(), f'{safe}_{safe_type}.docx')
    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)

    # Print verification summary
    para_count = len(doc.paragraphs)
    file_size = os.path.getsize(output_path)
    print(f"[OK] Document saved to: {output_path}")
    print(f"     Paragraphs: {para_count}")
    print(f"     File size: {file_size:,} bytes")

    return output_path


# ── JSON Config Loader ─────────────────────────────────────────────────────

def load_config_from_json(path):
    """Load a gongwen config dict from a JSON file.

    The JSON file must contain all required fields. Example:
    {
        "authority": "XX省人民政府办公厅",
        "document_type": "通知",
        "document_number": "X政发〔2024〕1号",
        ...
    }

    Args:
        path: Path to JSON config file.

    Returns:
        dict: Parsed configuration.
    """
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


# ── Template Library ──────────────────────────────────────────────────────

# Built-in template directory (relative to this script)
_BUILTIN_TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), '..', 'templates')

# User-settable override directory (set via CLI --template-dir or env var)
_USER_TEMPLATE_DIR = os.environ.get('GONGWEN_TEMPLATE_DIR', None)


def set_user_template_dir(path):
    """Set a user-defined template directory that overrides built-in templates.

    When set, load_template() searches this directory first, falling back
    to built-in templates only if the template_id is not found here.

    Args:
        path: Directory path containing template JSON files.
    """
    global _USER_TEMPLATE_DIR
    _USER_TEMPLATE_DIR = os.path.abspath(path)


def _find_template_file(template_id):
    """Find a template file by ID, respecting user override.

    Args:
        template_id: Template identifier (with or without .json suffix).

    Returns:
        (path, source_label) or (None, None).
        source_label is 'user', 'builtin', or None.
    """
    template_id = template_id.replace('.json', '')
    fname = f'{template_id}.json'

    # 1) User template dir first (override)
    if _USER_TEMPLATE_DIR:
        path = os.path.join(_USER_TEMPLATE_DIR, fname)
        if os.path.isfile(path):
            return path, 'user'

    # 2) Built-in
    path = os.path.join(_BUILTIN_TEMPLATE_DIR, fname)
    if os.path.isfile(path):
        return path, 'builtin'

    return None, None


def load_template(template_id):
    """Load a template JSON by ID.

    Search order:
      1. User custom template dir (set via --template-dir or GONGWEN_TEMPLATE_DIR)
      2. Built-in templates/

    Args:
        template_id: Template identifier (with or without .json suffix).

    Returns:
        (dict, str): (template_data, source_label) or (None, None).
        source_label is 'user' or 'builtin', useful for transparency.
    """
    path, source = _find_template_file(template_id)
    if path is None:
        return None, None
    with open(path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    return data, source


def load_template_from_file(file_path):
    """Load a template from an arbitrary file path.

    This is the mechanism for --template-file: users point to any
    .json file on disk, no directory scanning needed.

    Args:
        file_path: Absolute or relative path to a template JSON file.

    Returns:
        dict or None: Template data, or None if file doesn't exist or
        isn't valid JSON.
    """
    if not os.path.isfile(file_path):
        return None
    with open(file_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def list_templates():
    """List all available templates (user + builtin, user first).

    Returns:
        list of dicts with template_id, name, description, document_type, source.
    """
    seen = set()
    templates = []

    def _scan(dir_path, source_label):
        if not os.path.isdir(dir_path):
            return
        for fname in sorted(os.listdir(dir_path)):
            if not fname.endswith('.json') or fname.startswith('_'):
                continue
            tid = fname[:-5]
            if tid in seen:
                continue
            seen.add(tid)
            path = os.path.join(dir_path, fname)
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
            except (json.JSONDecodeError, UnicodeDecodeError):
                continue
            templates.append({
                'template_id': data.get('template_id', tid),
                'name': data.get('name', tid),
                'description': data.get('description', ''),
                'document_type': data.get('document_type', ''),
                'source': source_label,
            })

    # User templates first (higher priority)
    if _USER_TEMPLATE_DIR:
        _scan(_USER_TEMPLATE_DIR, 'user')
    # Built-in
    _scan(_BUILTIN_TEMPLATE_DIR, 'builtin')

    return templates


# ── CLI Entry Point ────────────────────────────────────────────────────────

def main():
    """CLI entry point: parse args and generate document."""
    parser = argparse.ArgumentParser(
        description='Generate a GB/T 9704-2012 compliant Chinese government official document.'
    )
    parser.add_argument(
        '--config', '-c',
        type=str,
        help='Path to JSON config file. If omitted, generates a sample 通知 document.'
    )
    parser.add_argument(
        '--output', '-o',
        type=str,
        default=None,
        help='Output .docx path (overrides config output_path if set).'
    )
    parser.add_argument(
        '--template', '-t',
        type=str,
        help='Template ID to print (e.g. notice-appointment). Shows template content and exits.'
    )
    parser.add_argument(
        '--template-file',
        type=str,
        help='Path to a single template JSON file to print (any location).'
    )
    parser.add_argument(
        '--template-dir', '-T',
        type=str,
        help='Directory of user custom templates (overrides built-in for same ID).'
    )
    parser.add_argument(
        '--list-templates', '-l',
        action='store_true',
        help='List all available templates (user + builtin) and exit.'
    )
    parser.add_argument(
        '--joint-sample',
        action='store_true',
        help='Generate sample joint-issuance document and exit.'
    )
    parser.add_argument(
        '--minutes-sample',
        action='store_true',
        help='Generate sample 纪要 (meeting minutes) document and exit.'
    )
    args = parser.parse_args()

    if args.joint_sample:
        generate_joint_sample()
        return

    if args.minutes_sample:
        generate_minutes_sample()
        return

    # --template-dir sets the override directory for this run
    if args.template_dir:
        set_user_template_dir(args.template_dir)

    if args.list_templates:
        templates = list_templates()
        if not templates:
            print("No templates found.")
            return
        print(f"\n{'Template ID':<30} {'Source':<8} {'文种':<6} Name")
        print("-" * 80)
        for t in templates:
            print(f"  {t['template_id']:<30} {t['source']:<8} {t['document_type']:<6} {t['name']}")
        return

    if args.template_file:
        tmpl = load_template_from_file(args.template_file)
        if not tmpl:
            print(f"[ERROR] Cannot load template file: {args.template_file}", file=sys.stderr)
            sys.exit(1)
        print(f"[INFO] Loaded from file: {args.template_file}")
        print(json.dumps(tmpl, ensure_ascii=False, indent=2))
        return

    if args.template:
        tmpl, source = load_template(args.template)
        if not tmpl:
            print(f"[ERROR] Template '{args.template}' not found.", file=sys.stderr)
            print(f"       Use --list-templates to see available templates.", file=sys.stderr)
            sys.exit(1)
        print(f"[INFO] Loaded template '{tmpl['name']}' (source: {source})")
        print(json.dumps(tmpl, ensure_ascii=False, indent=2))
        return

    if args.config:
        config = load_config_from_json(args.config)
        if args.output:
            config['output_path'] = args.output
        generate_gongwen(config)
    else:
        print("No config provided. Generating sample document...")
        generate_sample()


def generate_sample():
    """Generate a sample 通知 (notice) document for testing purposes."""
    config = {
        "authority": "XX省人民政府办公厅",
        "document_type": "通知",
        "document_number": "X政发〔2024〕1号",
        "title": "关于做好2024年元旦春节期间有关工作的通知",
        "main_recipient": "各市、县人民政府，省政府各部门、各直属机构：",
        "body_sections": [
            {
                "level": 0,
                "text": "一、提高思想认识，切实增强工作责任感"
            },
            {
                "level": 1,
                "text": "（一）充分认识做好元旦春节期间工作的重要性"
            },
            {
                "level": None,
                "text": "2024年是中华人民共和国成立75周年。做好元旦春节期间各项工作，确保人民群众度过欢乐祥和的节日，对于深入贯彻落实党中央、国务院决策部署，维护社会大局稳定，具有十分重要的意义。"
            },
            {
                "level": 1,
                "text": "（二）牢固树立以人民为中心的发展思想"
            },
            {
                "level": None,
                "text": "各地区各部门要始终把人民放在心中最高位置，以高度的政治责任感和历史使命感，扎实做好节日期间各项服务保障工作，不断增强人民群众的获得感、幸福感、安全感。"
            },
            {
                "level": 0,
                "text": "二、聚焦重点领域，扎实做好各项工作"
            },
            {
                "level": 1,
                "text": "（一）全力保障节日市场供应"
            },
            {
                "level": None,
                "text": "要密切监测重要民生商品市场供应和价格变化，加强产销对接和储备调节，确保节日期间市场供应充足、价格平稳。畅通物流配送渠道，保障民生物资运输畅通。"
            },
            {
                "level": 1,
                "text": "（二）切实做好群众出行保障"
            },
            {
                "level": None,
                "text": "要科学安排运力，优化运输组织，做好春运各项工作。加强交通拥堵路段疏导管控，及时发布出行信息提示。完善应急预案，妥善应对恶劣天气等突发情况。"
            },
            {
                "level": 1,
                "text": "（三）坚决守住安全生产底线"
            },
            {
                "level": None,
                "text": "要深入开展安全生产隐患排查整治，突出抓好矿山、危险化学品、交通运输、建筑施工、城镇燃气等重点行业领域安全监管。加强消防安全管理，严防火灾事故发生。"
            },
            {
                "level": 0,
                "text": "三、加强组织领导，确保各项工作落实到位"
            },
            {
                "level": None,
                "text": "各地区各部门要高度重视元旦春节期间各项工作，切实加强组织领导，周密安排部署，层层压实责任。主要负责同志要亲自抓、负总责，确保各项工作落到实处。"
            },
            {
                "level": None,
                "text": "要严格执行值班值守制度，遇有重大突发事件或重要紧急情况，要第一时间请示报告并采取有效措施妥善处置。要强化督查检查，对工作落实不力的严肃追责问责。"
            },
            {
                "level": None,
                "text": "各地区各部门要在2024年1月5日前将节日期间工作安排情况报送省政府办公厅。"
            },
        ],
        "attachments": [
            "重点任务分工表",
            "值班值守工作安排表",
        ],
        "issuing_org": "XX省人民政府办公厅",
        "date": "2024年12月20日",
        "cc_organs": [
            "省委办公厅",
            "省人大常委会办公厅",
            "省政协办公厅",
            "省纪委监委机关",
        ],
        "printing_org": "XX省人民政府办公厅",
        "printing_date": "2024年12月20日印发",
        "copies": "200",
    }

    output_dir = os.path.join(os.getcwd(), 'output_gongwen')
    os.makedirs(output_dir, exist_ok=True)
    config['output_path'] = os.path.join(output_dir, 'XX省人民政府办公厅_通知_元旦春节工作.docx')

    try:
        generate_gongwen(config)
        print(f"\nSample document generated successfully.")
        print(f"Open the file in Word to verify formatting.")
        print(f"Note: If fonts like '方正小标宋简体' are not installed, fallback fonts will be used.")
    except Exception as e:
        print(f"[ERROR] Failed to generate sample document: {e}", file=sys.stderr)
        sys.exit(1)


def generate_joint_sample():
    """Generate a sample joint-issuance document for testing."""
    config = {
        "joint_authorities": ["XX省教育厅", "XX省人力资源和社会保障厅"],
        "document_type": "通知",
        "document_number": "X教人〔2024〕15号",
        "title": "关于做好2024年全省教育系统职称评审工作的通知",
        "main_recipient": "各高等学校：",
        "body_sections": [
            {"level": 0, "text": "一、评审范围"},
            {"level": None, "text": "全省各类高等学校在职专业技术人员均可申报。"},
            {"level": 0, "text": "二、申报条件"},
            {"level": 1, "text": "（一）基本条件"},
            {"level": None, "text": "遵守国家法律法规，具有良好的职业道德和敬业精神。"},
            {"level": 0, "text": "三、申报材料"},
            {"level": None, "text": "（一）《专业技术职务任职资格评审表》一式两份。"},
            {"level": None, "text": "（二）学历证书、职称证书复印件。"},
        ],
        "joint_issuing_orgs": ["XX省教育厅", "XX省人力资源和社会保障厅"],
        "date": "2024年6月15日",
    }
    output_dir = os.path.join(os.getcwd(), 'output_gongwen')
    os.makedirs(output_dir, exist_ok=True)
    config['output_path'] = os.path.join(output_dir, '联合行文_职称评审通知.docx')
    generate_gongwen(config)


def generate_minutes_sample():
    """Generate a sample 纪要 (meeting minutes) document for testing."""
    config = {
        "authority": "XX省人民政府办公厅",
        "document_type": "纪要",
        "document_number": "",
        "title": "XX省人民政府常务会议纪要",
        "main_recipient": None,
        "meeting_info": {
            "time": "2024年12月20日上午9时",
            "location": "省政府会议室",
            "chair": "XX省长",
            "attendees": ["副省长XXX", "副省长XXX", "秘书长XXX"],
            "recorder": "XXX",
        },
        "body_sections": [
            {"level": 0, "text": "一、研究全省经济工作"},
            {"level": None, "text": "会议听取了省发展改革委关于当前经济运行情况的汇报。会议认为，今年以来全省经济运行总体平稳、稳中有进，但面临的形势依然复杂严峻。"},
            {"level": 0, "text": "二、研究安全生产工作"},
            {"level": None, "text": "会议听取了省应急管理厅关于安全生产情况的汇报。会议强调，要牢固树立安全发展理念，坚决守住安全生产底线。"},
            {"level": 0, "text": "三、议定事项"},
            {"level": None, "text": "（一）原则通过《关于进一步促进全省经济稳中向好的若干措施》，由省发展改革委根据会议意见修改完善后按程序印发。"},
            {"level": None, "text": "（二）同意成立省安全生产督导检查组，由省应急管理厅牵头，近期开展全省安全生产大检查。"},
        ],
        "issuing_org": "XX省人民政府办公厅",
        "date": "2024年12月22日",
        "printing_org": "XX省人民政府办公厅",
        "printing_date": "2024年12月22日印发",
    }
    output_dir = os.path.join(os.getcwd(), 'output_gongwen')
    os.makedirs(output_dir, exist_ok=True)
    config['output_path'] = os.path.join(output_dir, '纪要_常务会议.docx')
    generate_gongwen(config)


if __name__ == '__main__':
    main()
