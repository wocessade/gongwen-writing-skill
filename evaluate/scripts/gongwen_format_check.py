#!/usr/bin/env python3
"""Programmatic GB/T 9704-2012 format compliance checker — semantic version.

Usage:
    python gongwen_format_check.py <path_to.docx>

Outputs a JSON report with per-category pass/fail results and an overall
determination. Returns exit code 0 on success, 1 on error.

Semantic checking: each paragraph is classified by type (title, body,
heading level-1, heading level-2, red header, document number, etc.)
and validated against the font/size/alignment/spacing rules for that type.
"""

import argparse
import json
import os
import re
import sys
from typing import Any, Optional
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ── helpers ────────────────────────────────────────────────────────────

def _emu_to_mm(emu: int) -> float:
    """Convert EMU to millimetres."""
    return emu / 36000


def _emu_to_cm(emu: int) -> float:
    """Convert EMU to centimetres."""
    return emu / 36000 / 10


def _relaxed_eq(actual: float, expected: float, tolerance_mm: float = 1.5) -> bool:
    return abs(actual - expected) <= tolerance_mm


def _get_para_text(para: Any) -> str:
    return para.text.strip()


def _is_centered(para: Any) -> bool:
    return para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER


def _is_right_aligned(para: Any) -> bool:
    return para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.RIGHT


def _get_run_font_name(run: Any, axis: str = 'eastAsia') -> Optional[str]:
    """Get font name from a run's rFonts element on a specific axis."""
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None:
        return None
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        return None
    return rFonts.get(qn(f'w:{axis}'))


def _get_run_font_size(run: Any) -> Optional[float]:
    """Get font size in pt from a run."""
    if run.font.size is not None:
        return run.font.size.pt
    return None


def _is_run_red(run: Any) -> bool:
    """Check if a run has red-colored text."""
    c = run.font.color
    if c and c.rgb:
        rgb_str = str(c.rgb).upper()
        return rgb_str in ('FF0000', 'RED', '255000')
    return False


# ── Paragraph Classification ──────────────────────────────────────────

# Patterns for heading detection
_L1_HEADING_RE = re.compile(r'^[一二三四五六七八九十百]+、')
_L2_HEADING_RE = re.compile(r'^（[一二三四五六七八九十百]+）')
_L3_HEADING_RE = re.compile(r'^\d+\.')
_L4_HEADING_RE = re.compile(r'^\(\d+\)')

_PARA_TYPE_ORDER = [
    'security',           # 密级
    'urgency',            # 紧急程度
    'red_header',         # 发文机关标志(红头)
    'document_number',    # 发文字号
    'signatory',          # 签发人
    'title',              # 标题
    'main_recipient',     # 主送机关
    'meeting_info',      # 纪要信息
    'heading_level1',     # 一级标题
    'heading_level2',     # 二级标题
    'heading_level3',     # 三级标题
    'body',               # 正文
    'attachment_label',   # 附件说明
    'signature',          # 发文机关署名
    'date',               # 成文日期
    'cc',                 # 抄送
    'printing_info',      # 印发机关/日期
    'page_number',        # 页码
    'unknown',            # 无法分类
]


def classify_paragraph(para: Any, idx: int) -> str:
    """Classify a paragraph into a semantic type based on content and formatting.

    Returns one of the PARA_TYPE_ORDER types.
    """
    text = _get_para_text(para)
    if not text:
        return 'unknown'

    pf = para.paragraph_format
    alignment = pf.alignment

    # --- Red header detection (red text, first 5 paragraphs) ---
    if idx < 5:
        for run in para.runs:
            if _is_run_red(run):
                return 'red_header'

    # --- Security classification (short, bold, 密/绝/秘) ---
    if len(text) < 10 and re.match(r'^[绝机秘内]密', text):
        return 'security'

    # --- Urgency (短, bold, 特急/加急) ---
    if text in ('特急', '加急') and len(para.runs) > 0:
        run = para.runs[0]
        if run.bold:
            return 'urgency'

    # --- Document number (居中, 发文字号格式, 前5段) ---
    if idx < 8 and _is_centered(para) and len(text) < 40:
        if re.search(r'[〔\[（].+[〕\]）]', text) and '号' in text:
            return 'document_number'

    # --- 签发人 ---
    if text.startswith('签发人：') or '签发人' in text[:12]:
        return 'signatory'

    # --- Title (居中, 较长的完整句子, 不含编号) ---
    if _is_centered(para) and len(text) >= 8:
        if not _L1_HEADING_RE.match(text) and not _L2_HEADING_RE.match(text):
            if '关于' in text or '的' in text or text.endswith('纪要'):
                return 'title'

    # --- Main recipient (左对齐, 以冒号结尾, <80 chars) ---
    if text.endswith('：') and len(text) < 80:
        return 'main_recipient'

    # --- Heading level 1: 一、二、三、... ---
    if _L1_HEADING_RE.match(text):
        return 'heading_level1'

    # --- Heading level 2: （一）（二）（三）... ---
    if _L2_HEADING_RE.match(text):
        return 'heading_level2'

    # --- Heading level 3: 1. 2. 3. ... ---
    if _L3_HEADING_RE.match(text):
        return 'heading_level3'

    # --- Heading level 4: (1) (2) (3) ... ---
    if _L4_HEADING_RE.match(text):
        return 'heading_level4'

    # --- Attachment label ---
    if text.startswith('附件') and ('：' in text or ':' in text):
        return 'attachment_label'

    # --- CC (抄送) ---
    if text.startswith('抄送') or text.startswith('抄报'):
        return 'cc'

    # --- Printing info (印发) ---
    if '印发' in text and '份' in text:
        return 'printing_info'
    if idx > 5 and '印发' in text:
        return 'printing_info'

    # --- Signature/date detection (右对齐, 较短) ---
    if _is_right_aligned(para) and len(text) < 40:
        if re.search(r'\d{4}年', text):
            return 'date'
        if len(text) < 25:
            return 'signature'

    # --- Meeting info (纪要格式) ---
    if re.match(r'^(时 间|地 点|主持人|参加人员|记 录|出席|列席)', text):
        return 'meeting_info'

    # --- Body (default for long text) ---
    if len(text) > 15:
        return 'body'

    return 'unknown'


# ── Expected Fonts per Type ───────────────────────────────────────────

# (expected_font_eastAsia, expected_size_pt, expected_bold, expected_alignment)
_TYPE_EXPECTATIONS = {
    'red_header': {
        'font': '方正小标宋简体',
        'size_min': 28, 'size_max': 80,
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.CENTER,
        'color_red': True,
        'critical': True,
    },
    'title': {
        'font': '方正小标宋简体',
        'size_min': 18, 'size_max': 28,
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.CENTER,
        'critical': True,
    },
    'document_number': {
        'font': '仿宋_GB2312',
        'size_min': 14, 'size_max': 18,
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.CENTER,
        'critical': True,
    },
    'main_recipient': {
        'font': '仿宋_GB2312',
        'size_min': 14, 'size_max': 18,
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'critical': True,
    },
    'heading_level1': {
        'font': '黑体',
        'size_min': 14, 'size_max': 18,
        'bold': True,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'critical': False,
    },
    'heading_level2': {
        'font': '楷体_GB2312',
        'size_min': 14, 'size_max': 18,
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'critical': False,
    },
    'heading_level3': {
        'font': '仿宋_GB2312',
        'size_min': 14, 'size_max': 18,
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'critical': False,
    },
    'heading_level4': {
        'font': '仿宋_GB2312',
        'size_min': 14, 'size_max': 18,
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'critical': False,
    },
    'body': {
        'font': '仿宋_GB2312',
        'size_min': 14, 'size_max': 18,
        'bold': False,
        'alignment': None,  # left-aligned with first-line indent
        'critical': False,
    },
    'signatory': {
        'font': '仿宋_GB2312',
        'size_min': 14, 'size_max': 18,
        'bold': False,
        'alignment': None,
        'critical': False,
    },
    'signature': {
        'font': '仿宋_GB2312',
        'size_min': 14, 'size_max': 18,
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.RIGHT,
        'critical': False,
    },
    'date': {
        'font': '仿宋_GB2312',
        'size_min': 14, 'size_max': 18,
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.RIGHT,
        'critical': False,
    },
    'attachment_label': {
        'font': '仿宋_GB2312',
        'size_min': 14, 'size_max': 18,
        'bold': False,
        'alignment': None,
        'critical': False,
    },
    'cc': {
        'font': '仿宋_GB2312',
        'size_min': 14, 'size_max': 18,
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'critical': False,
    },
    'printing_info': {
        'font': '仿宋_GB2312',
        'size_min': 14, 'size_max': 18,
        'bold': False,
        'alignment': None,
        'critical': False,
    },
    'security': {
        'font': '黑体',
        'size_min': 14, 'size_max': 18,
        'bold': True,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'critical': True,
    },
    'urgency': {
        'font': '黑体',
        'size_min': 14, 'size_max': 18,
        'bold': True,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'critical': True,
    },
    'meeting_info': {
        'font': '仿宋_GB2312',
        'size_min': 14, 'size_max': 18,
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'critical': False,
    },
    'unknown': {
        'font': None,
        'size_min': None, 'size_max': None,
        'bold': None,
        'alignment': None,
        'critical': False,
    },
}


# ── Check Functions ───────────────────────────────────────────────────

_MARGIN_ATTRS = {
    "top_margin": ("top_margin", 37.0),
    "bottom_margin": ("bottom_margin", 35.0),
    "left_margin": ("left_margin", 28.0),
    "right_margin": ("right_margin", 26.0),
}


def check_page_setup(doc: Any) -> dict:
    """Check page margins against GB/T 9704-2012 standard values."""
    section = doc.sections[0] if doc.sections else None
    if section is None:
        return {"pass": False, "error": "No sections found in document"}

    checks: dict[str, dict] = {}
    all_pass = True

    for attr_name, (check_name, expected_mm) in _MARGIN_ATTRS.items():
        emu = getattr(section, attr_name, None)
        if emu is None:
            checks[check_name] = {
                "expected": f"{expected_mm}mm",
                "actual": "not set",
                "pass": False,
            }
            all_pass = False
            continue
        actual_mm = round(_emu_to_mm(emu), 1)
        ok = _relaxed_eq(actual_mm, expected_mm)
        if not ok:
            all_pass = False
        checks[check_name] = {
            "expected": f"{expected_mm}mm",
            "actual": f"{actual_mm}mm",
            "pass": ok,
        }

    return {"pass": all_pass, "checks": checks}


def check_semantic_fonts(doc: Any) -> dict:
    """Semantic font check: classify each paragraph and validate per-type.

    Returns pass/fail and a list of per-paragraph issues.
    """
    issues: list[dict] = []
    classification_summary: dict[str, int] = {}
    font_accept_fallbacks = {
        '方正小标宋简体': {'方正小标宋简体', '方正小标宋', '宋体', 'SimSun'},
        '仿宋_GB2312': {'仿宋_GB2312', '仿宋', 'FangSong', '宋体', 'SimSun'},
        '楷体_GB2312': {'楷体_GB2312', '楷体', 'KaiTi', '宋体', 'SimSun'},
        '黑体': {'黑体', 'SimHei', '宋体', 'SimSun'},
        '宋体': {'宋体', 'SimSun'},
    }

    for idx, para in enumerate(doc.paragraphs):
        text = _get_para_text(para)
        if not text:
            continue

        ptype = classify_paragraph(para, idx)
        classification_summary[ptype] = classification_summary.get(ptype, 0) + 1

        exp = _TYPE_EXPECTATIONS.get(ptype)
        if exp is None or ptype == 'unknown':
            continue

        # Check each run in the paragraph
        para_issues = []
        for run_idx, run in enumerate(para.runs):
            if not run.text.strip():
                continue

            # Font name check
            font_name = _get_run_font_name(run, 'eastAsia')
            if font_name and exp['font']:
                allowed = font_accept_fallbacks.get(exp['font'], {exp['font']})
                if font_name not in allowed:
                    para_issues.append(
                        f"Run {run_idx}: font '{font_name}' (expected '{exp['font']}')"
                    )

            # Font size check
            font_size = _get_run_font_size(run)
            if font_size is not None and exp['size_min'] is not None:
                if font_size < exp['size_min'] or font_size > exp['size_max']:
                    para_issues.append(
                        f"Run {run_idx}: size {font_size:.0f}pt (expected "
                        f"{exp['size_min']}-{exp['size_max']}pt)"
                    )

        # Alignment check (from paragraph format, not per-run)
        if exp.get('alignment') is not None and para_issues:
            actual_align = para.paragraph_format.alignment
            if actual_align != exp['alignment'] and actual_align is not None:
                para_issues.append(
                    f"alignment {actual_align} (expected {exp['alignment']})"
                )

        # Red color check for red header
        if ptype == 'red_header':
            has_red = any(_is_run_red(run) for run in para.runs)
            if not has_red:
                para_issues.append("red header text is not red (expected RGB 255,0,0)")

        if para_issues:
            issues.append({
                "paragraph": idx,
                "type": ptype,
                "text_preview": text[:50],
                "issues": para_issues,
            })

    # Determine pass/fail: critical type errors are blocking
    critical_fail = False
    for issue in issues:
        ptype = issue['type']
        exp = _TYPE_EXPECTATIONS.get(ptype, {})
        if exp.get('critical', False):
            critical_fail = True

    return {
        "pass": not critical_fail,
        "issues": issues,
        "classification": classification_summary,
        "total_paragraphs": len(doc.paragraphs),
        "critical_fail": critical_fail,
    }


def check_line_spacing(doc: Any) -> dict:
    """Check body paragraphs for fixed 28pt line spacing."""
    issues: list[str] = []

    for idx, para in enumerate(doc.paragraphs):
        text = _get_para_text(para)
        if not text:
            continue

        ptype = classify_paragraph(para, idx)

        # Only check body text, headings, and meeting info for line spacing
        if ptype not in ('body', 'heading_level1', 'heading_level2',
                         'heading_level3', 'heading_level4', 'meeting_info',
                         'main_recipient', 'signature', 'date', 'attachment_label',
                         'signatory'):
            continue

        pf = para.paragraph_format
        if pf.line_spacing is None:
            continue

        try:
            spacing_val = pf.line_spacing
            if hasattr(spacing_val, "pt"):
                pt = spacing_val.pt
                if abs(pt - 28) > 4:
                    issues.append(
                        f"Paragraph {idx} ({ptype}): line spacing is {pt:.0f}pt "
                        f"(expected ~28pt) — \"{text[:40]}\""
                    )
            elif isinstance(spacing_val, float):
                issues.append(
                    f"Paragraph {idx} ({ptype}): line spacing is a multiple "
                    f"({spacing_val:.1f}) instead of fixed 28pt — \"{text[:40]}\""
                )
        except (TypeError, AttributeError):
            pass

    return {"pass": len(issues) == 0, "issues": issues}


def check_indent(doc: Any) -> dict:
    """Check body paragraphs for ~2-character first-line indent."""
    issues: list[str] = []

    for idx, para in enumerate(doc.paragraphs):
        text = _get_para_text(para)
        if not text:
            continue

        ptype = classify_paragraph(para, idx)

        # Only body text needs first-line indent
        if ptype not in ('body', 'meeting_info'):
            continue

        pf = para.paragraph_format
        # Skip if centered or right-aligned
        if pf.alignment in (WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT):
            continue

        fi = pf.first_line_indent
        if fi is None or fi == 0:
            issues.append(
                f"Paragraph {idx} ({ptype}): no first-line indent — \"{text[:40]}\""
            )
        else:
            try:
                indent_mm = _emu_to_mm(fi)
                if indent_mm < 5:
                    issues.append(
                        f"Paragraph {idx}: first-line indent too small "
                        f"({indent_mm:.0f}mm, expected ~11mm / 32pt) — \"{text[:40]}\""
                    )
            except Exception:
                pass

    return {"pass": len(issues) == 0, "issues": issues}


def check_red_header(doc: Any) -> dict:
    """Check for red header text and red divider line."""
    found_red_text = False
    found_divider = False
    red_para_idx = None

    for idx, para in enumerate(doc.paragraphs[:15]):
        for run in para.runs:
            if _is_run_red(run):
                found_red_text = True
                red_para_idx = idx
                break
        if found_red_text:
            break

    # Check for red divider line (pBdr with red color) within first 15 paragraphs
    for idx, para in enumerate(doc.paragraphs[:15]):
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is not None:
                for border in pBdr:
                    color = border.get(qn("w:color"))
                    if color and color.upper() == "FF0000":
                        found_divider = True
                        break

    issues = []
    if not found_red_text:
        issues.append("No red-colored text found in first 15 paragraphs (expected red header)")
    else:
        issues.append(f"Red header text found at paragraph {red_para_idx}")
    if not found_divider:
        issues.append("No red divider line (pBdr with red color) detected below the header")
    else:
        issues.append("Red divider line detected")

    return {"pass": found_red_text and found_divider, "issues": issues}


def check_page_number(doc: Any) -> dict:
    """Check footer for page number field with '- N -' format."""
    found_page_field = False

    for section in doc.sections:
        footers_to_check = [section.footer]
        try:
            footers_to_check.append(section.even_page_footer)
        except Exception:
            pass
        for footer in footers_to_check:
            if footer is None:
                continue
            for instr in footer._element.iter(qn("w:instrText")):
                if instr.text and "PAGE" in instr.text.upper():
                    found_page_field = True
                    break
            for para in footer.paragraphs:
                text = para.text.strip()
                if ("-" in text or "—" in text) and any(ch.isdigit() for ch in text):
                    found_page_field = True

    issues = []
    if not found_page_field:
        issues.append("No page number field or '- N -' pattern found in footer")

    return {"pass": found_page_field, "issues": issues}


# ── aggregate ─────────────────────────────────────────────────────────

def check_all(doc_path: str) -> dict:
    """Run all checks on *doc_path* and return a JSON-serialisable report."""
    from docx import Document

    if not os.path.isfile(doc_path):
        raise FileNotFoundError(f"Document not found: {doc_path}")

    doc = Document(doc_path)

    page_setup = check_page_setup(doc)
    semantic_fonts = check_semantic_fonts(doc)
    line_spacing = check_line_spacing(doc)
    indent = check_indent(doc)
    red_header = check_red_header(doc)
    page_numbers = check_page_number(doc)

    # Count issues by severity
    critical = 0
    major = 0
    minor = 0

    if not page_setup.get("pass", False):
        critical += 1
    if not red_header.get("pass", False):
        critical += 1
    if semantic_fonts.get("critical_fail", False):
        critical += 1
    if not line_spacing.get("pass", False):
        major += 1
    if not indent.get("pass", False):
        major += 1
    if not page_numbers.get("pass", False):
        minor += 1

    # Font issues classified by severity of the paragraph type
    if semantic_fonts.get("issues"):
        for issue in semantic_fonts["issues"]:
            ptype = issue['type']
            exp = _TYPE_EXPECTATIONS.get(ptype, {})
            if exp.get('critical', False):
                critical += 1  # already counted above, additional detail
            elif ptype in ('body', 'heading_level1', 'heading_level2',
                           'heading_level3', 'heading_level4'):
                major += 1

    overall_pass = (critical == 0 and major <= 2 and minor <= 3)

    return {
        "file": os.path.abspath(doc_path),
        "page_setup": page_setup,
        "semantic_fonts": semantic_fonts,
        "line_spacing": line_spacing,
        "indent": indent,
        "red_header": red_header,
        "page_numbers": page_numbers,
        "overall": "pass" if overall_pass else "fail",
        "critical_issues": critical,
        "major_issues": major,
        "minor_issues": minor,
    }


# ── CLI ───────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="GB/T 9704-2012 format compliance checker for .docx files"
    )
    parser.add_argument("doc_path", help="Path to a .docx file")
    args = parser.parse_args()

    try:
        report = check_all(args.doc_path)
    except Exception as exc:
        print(json.dumps({"error": str(exc)}, ensure_ascii=False, indent=2))
        sys.exit(1)

    print(json.dumps(report, ensure_ascii=False, indent=2))
    sys.exit(0 if report["overall"] == "pass" else 1)


if __name__ == "__main__":
    main()
